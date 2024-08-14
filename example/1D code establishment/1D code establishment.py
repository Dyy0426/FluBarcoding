from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re


def replace_bases_and_format(input_file, output_file):
    # 确定后缀以选择处理方式
    file_suffix = input_file.split('.')[-1].lower()
    text = ""

    # 读取文件内容
    if file_suffix in ['doc', 'docx']:
        doc = Document(input_file)
        text = '\n'.join([p.text for p in doc.paragraphs])
    elif file_suffix == 'txt':
        with open(input_file, 'r') as file:
            text = file.read()
    else:
        raise ValueError("Unsupported file format")

    # 替换碱基并记录替换位置和碱基类型
    base_replacements = {
        'A': ('|', 'Showcard Gothic', 14, RGBColor(0x80, 0x76, 0xA3)),  # 四号字体大约对应12磅
        'T': ('|', 'Showcard Gothic', 14, RGBColor(0xFF, 0x00, 0x00)),
        'G': ('|', 'Showcard Gothic', 16, RGBColor(0x56, 0x98, 0xC3)),  # 三号字体大约对应16磅
        'C': ('|', 'Showcard Gothic', 16, RGBColor(0x00, 0xB0, 0x50)),
    }

    # 创建新的Word文档
    doc = Document()
    # 设置整体样式
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Showcard Gothic'

    # 将文本添加到文档
    for line in text.split('\n'):
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        for char in line:
            run = p.add_run(char)
            if char in base_replacements:
                replacement, font_name, font_size, color = base_replacements[char]
                run.text = replacement
                run.font.name = font_name
                run.font.size = Pt(font_size)
                run.font.color.rgb = color

    # 保存新的Word文档
    doc.save(output_file)


# 调用函数处理示例文件
input_file = 'input.txt'  # 更换为实际文件路径
output_file = 'output.docx'  # 更换为所需的输出文件路径
replace_bases_and_format(input_file, output_file)