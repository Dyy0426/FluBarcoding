from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def replace_bases_and_format(input_file, output_file):
    # Determine the file suffix to choose the processing method
    file_suffix = input_file.split('.')[-1].lower()
    text = ""

    # Read the file content
    if file_suffix in ['doc', 'docx']:
        doc = Document(input_file)
        text = '\n'.join([p.text for p in doc.paragraphs])
    elif file_suffix == 'txt':
        with open(input_file, 'r') as file:
            text = file.read()
    else:
        raise ValueError("Unsupported file format")

    # Replace bases and record the replacement positions and base types
    base_replacements = {
        'A': ('|', 'Showcard Gothic', 14, RGBColor(0x80, 0x76, 0xA3)),  # Size 14 font roughly corresponds to 12pt
        'T': ('|', 'Showcard Gothic', 14, RGBColor(0xFF, 0x00, 0x00)),
        'G': ('|', 'Showcard Gothic', 16, RGBColor(0x56, 0x98, 0xC3)),  # Size 16 font roughly corresponds to 16pt
        'C': ('|', 'Showcard Gothic', 16, RGBColor(0x00, 0xB0, 0x50)),
    }

    # Create a new Word document
    doc = Document()
    # Set the overall style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Showcard Gothic'

    # Add text to the document
    for line in text.split('\n'):
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        for char in line:
            run = p.add_run(char)
            if char in base_replacements:
                replacement, font_name, font_size, color = base_replacements[char]
                run.text = replacement
                run.font.name = font_name
                run.font.size = Pt(font_size)
                run.font.color.rgb = color

    # Save the new Word document
    doc.save(output_file)

# Call the function to process an example file
input_file = 'input.docx'  # Replace with the actual file path
output_file = 'output.docx'  # Replace with the desired output file path
replace_bases_and_format(input_file, output_file)
